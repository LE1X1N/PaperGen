from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.oxml import OxmlElement


class StyleController:
    def __init__(self):
        self.COVER_SCHOOL_STYLE = "论文_封面_学校"
        self.COVER_HEAD_STYLE = "论文_封面_标题"
        self.COVER_TITLE_STYLE = "论文_封面_题目"
        self.COVER_TABLE_LEFT_COL_STYLE = "论文_封面_表左"
        self.COVER_TABLE_RIGHT_COL_STYLE = "论文_封面_表右"

        self.KEYWORD_STYLE = "论文_关键字"

        self.TOC_HEAD_STYLE = "论文_目录"
        self.TOC_STYLE_1 = "论文_目录 1"
        self.TOC_STYLE_2 = "论文_目录 2"
        self.TOC_STYLE_3 = "论文_目录 3"
        
        self.NORMAL_STYLE = "论文_正文"
        
        self.HEADING_STYLE_1 = "论文_标题 1"
        self.HEADING_STYLE_2 = "论文_标题 2"
        self.HEADING_STYLE_3 = "论文_标题 3"

    def init_doc_style(self, doc: Document):
        self._modify_cover_school_style(doc)
        self._modify_cover_head_style(doc)
        self._modify_cover_title_style(doc)
        self._modify_cover_table_left_col_style(doc)
        self._modify_cover_table_right_col_style(doc)

        self._modify_keyword_style(doc)

        self._modify_toc_head_styles(doc)
        self._modify_toc_l1_styles(doc)
        self._modify_toc_l2_styles(doc)
        self._modify_toc_l3_styles(doc)

        self._modify_normal_style(doc)
        
        self._modify_heading_l1_styles(doc)
        self._modify_heading_l2_styles(doc)
        self._modify_heading_l3_styles(doc)

    
    def _modify_cover_school_style(self, doc: Document):
        style = doc.styles.add_style(self.COVER_SCHOOL_STYLE, WD_STYLE_TYPE.PARAGRAPH)

        style.font.name = "Times New Roman"
        style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        style.font.size = Pt(24)
        style.font.bold = True
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_before = Pt(18)
        style.paragraph_format.space_after = Pt(18)

    def _modify_cover_head_style(self, doc: Document):
        style = doc.styles.add_style(self.COVER_HEAD_STYLE, WD_STYLE_TYPE.PARAGRAPH)

        style.font.name = "Times New Roman"
        style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        style.font.size = Pt(36)
        style.font.bold = True
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_before = Pt(30)
        style.paragraph_format.space_after = Pt(24)

    def _modify_cover_title_style(self, doc: Document):
        style = doc.styles.add_style(self.COVER_TITLE_STYLE, WD_STYLE_TYPE.PARAGRAPH)

        style.font.name = "Times New Roman"
        style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(18)
        style.font.bold = True
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_before = Pt(120)
        style.paragraph_format.space_after = Pt(90)


    def _modify_cover_table_left_col_style(self, doc: Document):
        style = doc.styles.add_style(self.COVER_TABLE_LEFT_COL_STYLE, WD_STYLE_TYPE.CHARACTER)

        style.font.name = "Times New Roman"
        style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(16)
        style.font.bold = False

    def _modify_cover_table_right_col_style(self, doc: Document):
        style = doc.styles.add_style(self.COVER_TABLE_RIGHT_COL_STYLE, WD_STYLE_TYPE.CHARACTER)

        style.font.name = "Times New Roman"
        style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(16)
        style.font.bold = False

    def _modify_keyword_style(self, doc: Document):
        style = doc.styles.add_style(self.KEYWORD_STYLE, WD_STYLE_TYPE.PARAGRAPH)
            
        # font style
        font_size = Pt(12)
        style.font.size = font_size
        style.font.bold = False
        style.font.name = "Times New Roman"
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        style.paragraph_format.line_spacing = 1.5

    def _modify_toc_head_styles(self, doc: Document):
        style = doc.styles.add_style(self.TOC_HEAD_STYLE, WD_STYLE_TYPE.PARAGRAPH)

        style.font.name = "Times New Roman"
        style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.font.size = Pt(16)
        style.font.bold = True
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.line_spacing = 1.5

    def _modify_toc_l1_styles(self, doc: Document):
        style = doc.styles.add_style(self.TOC_STYLE_1, WD_STYLE_TYPE.PARAGRAPH)

        style.font.size = Pt(12)  # 12磅（小四）
        style.font.bold = True
        style.font.name = "Times New Roman"
        style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    def _modify_toc_l2_styles(self, doc: Document):
        style = doc.styles.add_style(self.TOC_STYLE_2, WD_STYLE_TYPE.PARAGRAPH)

        style.font.size = Pt(12)  # 12磅（小四）
        style.font.bold = False
        style.font.name = "Times New Roman"
        style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.left_indent = Pt(20)

    def _modify_toc_l3_styles(self, doc: Document):
        style = doc.styles.add_style(self.TOC_STYLE_3, WD_STYLE_TYPE.PARAGRAPH)

        style.font.size = Pt(12)  # 12磅（小四）
        style.font.bold = False
        style.font.name = "Times New Roman"
        style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.left_indent = Pt(40)

    def _modify_normal_style(self, doc: Document):
        style = doc.styles.add_style(self.NORMAL_STYLE, WD_STYLE_TYPE.PARAGRAPH)
            
        # font style
        font_size = Pt(12)
        style.font.size = font_size
        style.font.bold = False
        style.font.name = "Times New Roman"
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # jusify alignment, first line indent 2 characters
        style.paragraph_format.first_line_indent = font_size * 2 
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        style.paragraph_format.line_spacing = 1.5

    def _modify_heading_l1_styles(self, doc: Document):
        style = doc.styles.add_style(self.HEADING_STYLE_1, WD_STYLE_TYPE.PARAGRAPH)

        # font
        style.font.size = Pt(16)
        style.font.bold = True
        style.font.name = "Times New Roman"
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # paragraph
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        style.paragraph_format.space_before = Pt(13)

    def _modify_heading_l2_styles(self, doc: Document):
        style = doc.styles.add_style(self.HEADING_STYLE_2, WD_STYLE_TYPE.PARAGRAPH)

        # font
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.name = "Times New Roman"
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # paragraph
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        style.paragraph_format.space_before = Pt(13)

    def _modify_heading_l3_styles(self, doc: Document):
        style = doc.styles.add_style(self.HEADING_STYLE_3, WD_STYLE_TYPE.PARAGRAPH)

        # font
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.name = "Times New Roman"
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # paragraph
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        style.paragraph_format.space_before = Pt(13)


    def modify_cover_table_style(self, table: Table):
        # table dimension
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            row.height = Cm(1.2)

        # text style
        for row in table.rows:
            # left col
            row.cells[0].width = Cm(3)
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            para0 = row.cells[0].paragraphs[0]
            para0.paragraph_format.space_before = Pt(0)
            para0.paragraph_format.space_after = Pt(0)
            para0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para0.runs[0]
            run.font.name = "Times New Roman"
            run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(16)
            run.font.bold = False

            # right col
            row.cells[1].width = Cm(5.5)
            row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            para1 = row.cells[1].paragraphs[0]
            para1.paragraph_format.space_before = Pt(0)
            para1.paragraph_format.space_after = Pt(0)
            para1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para1.runs[0]
            run.font.name = "Times New Roman"
            run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(16)
            run.font.bold = False
        
        # border
        for row in table.rows:
            right_cell = row.cells[1]
            tc = right_cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            border = OxmlElement(f'w:bottom')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')  
            tcPr.append(border)


    def modify_main_body_table_style(self, table: Table):
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # col width
        num_cols = len(table.columns)
        col_maxlens = [0] * num_cols
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                cell_text = "".join([para.text for para in cell.paragraphs])
                col_maxlens[col_idx] = max(col_maxlens[col_idx], len(cell_text))

        total_len = sum(col_maxlens) if sum(col_maxlens) > 0 else num_cols
        total_width = Pt(400)
        for col_idx, col in enumerate(table.columns):
            col_width = total_width * (col_maxlens[col_idx] / total_len)
            for cell in col.cells:
                cell.width = col_width

        # alignment
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # text style
        def _set_cell_text_style(cell: _Cell , is_header: bool = False):
            for para in cell.paragraphs:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0  

                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

                    if is_header:
                        run.font.size = Pt(12)
                        run.font.bold = True
                    else:
                        run.font.size = Pt(11)
                        run.font.bold = False

        for cell in table.rows[0].cells:
            _set_cell_text_style(cell, is_header=True)
        for row in table.rows[1:]:
            for cell in row.cells:
                _set_cell_text_style(cell, is_header=False)

        # border
        # header row
        for cell in table.rows[0].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            border = OxmlElement(f'w:top')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')  
            tcPr.append(border)
            
            border = OxmlElement(f'w:bottom')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')  
            tcPr.append(border)

        # bottom row
        for cell in table.rows[-1].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            border = OxmlElement(f'w:bottom')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')  
            tcPr.append(border)
