from docx.document import Document
from datetime import datetime
from src.service.doc.style_controller import StyleController
from typing import List

class DocComposer:
    def __init__(self, doc: Document):
        self.doc = doc
        self.style_controller = StyleController()

        # init styles of doc
        self.style_controller.init_doc_style(self.doc)

    def compose_all(self, title, abstract=None, structure=None, main_body=None, tables=None):
        self.compose_cover(title)
        if abstract is not None:
            self.compose_abstract(abstract)
        if structure is not None:
            self.compose_toc(structure)
        if main_body is not None:
            self.compose_main_body(structure, main_body, tables)
        

    def compose_cover(self, title: str):
        self.doc.add_paragraph("XXXX大学", style=self.style_controller.COVER_SCHOOL_STYLE)
        self.doc.add_paragraph("本科生毕业论文", style=self.style_controller.COVER_HEAD_STYLE)
        self.doc.add_paragraph(title, style=self.style_controller.COVER_TITLE_STYLE)
        
        table = self.doc.add_table(rows=0, cols=2)
        info = (
            ("姓        名:", "XXX"),
            ("学        号:", "123456"),
            ("院        系:", "计算机学院"),
            ("专        业:", "软件工程"),
            ("指导教师:", "XXX"),
            ("完成日期:", datetime.now().strftime("%Y年%m月%d日"))
        )

        for k, v in info:
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].add_run(k)
            row_cells[1].paragraphs[0].add_run(v)

        self.style_controller.modify_cover_table_style(table) 
        self.doc.add_page_break()


    def compose_abstract(self, abstract: dict):
        # generate abstract (both Chinese and English)
        self.doc.add_paragraph("摘要", style=self.style_controller.TOC_HEAD_STYLE)
        self.doc.add_paragraph(abstract["abstract_cn"], style=self.style_controller.NORMAL_STYLE)
        self.doc.add_paragraph(f"关键字：{abstract['keyword_cn']}", style=self.style_controller.KEYWORD_STYLE)
        self.doc.add_page_break()

        self.doc.add_paragraph("ABSTRACT", style=self.style_controller.TOC_HEAD_STYLE)
        self.doc.add_paragraph(abstract["abstract_en"], style=self.style_controller.NORMAL_STYLE)
        self.doc.add_paragraph(f"Keywords: {abstract['keyword_en']}", style=self.style_controller.KEYWORD_STYLE)
        self.doc.add_page_break()


    def compose_toc(self, structure: dict):
        # generate table of contents in a docx
        self.doc.add_paragraph("目录", style=self.style_controller.TOC_HEAD_STYLE)

        for chapter in structure["chapters"]:
            # 1-level
            self.doc.add_paragraph(chapter["title"], style=self.style_controller.TOC_STYLE_1)

            # 2-level
            if "sections" in chapter:
                for section in chapter["sections"]:
                    self.doc.add_paragraph(section["title"], style=self.style_controller.TOC_STYLE_2)

                    # 3-level 
                    if "subsections" in section:
                        for subsection in section["subsections"]:
                            self.doc.add_paragraph(subsection["title"], style=self.style_controller.TOC_STYLE_3) 
        self.doc.add_page_break() 


    def compose_main_body(self, structure: dict, main_body: dict, tables: dict=None):
        
        for chapter in structure["chapters"]:
            # 1-level Handling
            self.doc.add_paragraph(chapter["title"], style=self.style_controller.HEADING_STYLE_1)

            if "sections" not in chapter:
                # Main Body     
                texts = main_body[chapter['title']]
                for text in texts:
                    self.doc.add_paragraph(text, style=self.style_controller.NORMAL_STYLE) 
                
                # compose table
                if chapter['title'] in tables:
                    self._compose_table(tables[chapter['title']])

            # 2-level
            else:
                for section in chapter["sections"]:
                    # 2-level Handling
                    self.doc.add_paragraph(section['title'], style=self.style_controller.HEADING_STYLE_2)

                    if "subsections" not in section:
                        # Main Body
                        texts = main_body[section['title']]
                        for text in texts:
                            self.doc.add_paragraph(text, style=self.style_controller.NORMAL_STYLE)

                        # compose table
                        if section['title'] in tables:
                            self._compose_table(tables[section['title']])

                    # 3-level 
                    else:
                        for subsection in section["subsections"]:
                            # 3-level Handling
                            self.doc.add_paragraph(subsection['title'], style=self.style_controller.HEADING_STYLE_3)
                            # Main Body
                            texts = main_body[subsection['title']]
                            for text in texts:
                                self.doc.add_paragraph(text, style=self.style_controller.NORMAL_STYLE)

                            # compose table
                            if subsection['title'] in tables:
                                self._compose_table(tables[subsection['title']])
                        
            self.doc.add_page_break() 

    def _compose_table(self, table_data: dict | List[dict]):
        """
            insert tables into a docx

            table: dict
                e.g.
                {
                    "id": "表1-1 学生成绩表"
                    "headers": ["学号", "姓名", "语文", "数学", "英语"], 
                    "rows": [
                        ["001", "张三", "90", "85", "92"],
                        ["002", "李四", "88", "91", "89"],
                        ["003", "王五", "95", "87", "93"]
                    ]
                }
        """
        if isinstance(table_data, dict):
            table_data = [table_data]
            
        for data in table_data:
            # add tatle title
            self.doc.add_paragraph(data['id'], style=self.style_controller.TABLE_TITLE_STYLE)

            # add table
            table = self.doc.add_table(rows=1, cols=len(data["headers"]))
            # header
            for idx, cell in enumerate(table.rows[0].cells):
                cell.text = data["headers"][idx]
            
            # row
            for row in data["rows"]:
                row_cells = table.add_row().cells
                for idx, cell in enumerate(row_cells):
                    cell.text = row[idx]

            self.style_controller.modify_main_body_table_style(table) 
