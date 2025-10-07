from docx.document import Document
import docx
import os
from datetime import datetime
from src.service.doc.style_controller import StyleController

style_controller = StyleController()


def init_doc(file_path: str=None):
    # initialize a doc object and styles
    if os.path.exists(file_path):
        doc = docx.Document(file_path)
        print(f"打开文档: {file_path}")
    else:
        doc = docx.Document()
        print(f"创建新文档: {file_path}")

    style_controller.init_doc_style(doc)
    return doc


def compose_cover(doc: Document, title: str):
    doc.add_paragraph("XXXX大学", style=style_controller.COVER_SCHOOL_STYLE)
    doc.add_paragraph("本科生毕业论文", style=style_controller.COVER_HEAD_STYLE)
    doc.add_paragraph(title, style=style_controller.COVER_TITLE_STYLE)
    
    table = doc.add_table(rows=0, cols=2)
    info = (
        ("姓名:", "XXX"),
        ("学号:", "123456"),
        ("院系:", "计算机学院"),
        ("专业:", "软件工程"),
        ("指导教师:", "XXX"),
        ("完成日期:", datetime.now().strftime("%Y年%m月%d日"))
    )

    for k, v in info:
        row_cells = table.add_row().cells
        row_cells[0].text = k
        row_cells[1].text = v

    doc.add_page_break()


def compose_toc(doc: Document, structure: dict):
    # generate table of contents in a docx
    doc.add_paragraph("目录", style=style_controller.TOC_HEAD_STYLE)

    for chapter in structure["chapters"]:
        # 1-level
        doc.add_paragraph(chapter["title"], style=style_controller.TOC_STYLE_1)

        # 2-level
        if "sections" in chapter:
            for section in chapter["sections"]:
                doc.add_paragraph(section["title"], style=style_controller.TOC_STYLE_2)

                # 3-level 
                if "subsections" in section:
                    for subsection in section["subsections"]:
                        doc.add_paragraph(subsection["title"], style=style_controller.TOC_STYLE_3) 
    doc.add_page_break() 


def compose_main_body(doc: Document, structure: dict, main_body: dict):
    
    for chapter in structure["chapters"]:
        # 1-level Handling
        doc.add_paragraph(chapter["title"], style=style_controller.HEADING_STYLE_1)

        if "sections" not in chapter:
            # Main Body     
            texts = main_body[chapter['title']]
            for text in texts:
                doc.add_paragraph(text, style=style_controller.NORMAL_STYLE) 

        # 2-level
        else:
            for section in chapter["sections"]:
                # 2-level Handling
                doc.add_paragraph(section['title'], style=style_controller.HEADING_STYLE_2)

                if "subsections" not in section:
                    # Main Body
                    texts = main_body[section['title']]
                    for text in texts:
                        doc.add_paragraph(text, style=style_controller.NORMAL_STYLE)

                # 3-level 
                else:
                    for subsection in section["subsections"]:
                        # 3-level Handling
                        doc.add_paragraph(subsection['title'], style=style_controller.HEADING_STYLE_3)
                        # Main Body
                        texts = main_body[subsection['title']]
                        for text in texts:
                            doc.add_paragraph(text, style=style_controller.NORMAL_STYLE)

        doc.add_page_break() 
