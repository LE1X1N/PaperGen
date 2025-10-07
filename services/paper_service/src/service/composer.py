from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os
from tqdm import tqdm
from src.service.gen_funcs import generate_section_text


TOC_STYLE_1 = "目录 1"
TOC_STYLE_2 = "目录 2"
TOC_STYLE_3 = "目录 3"
NORMAL_STYLE = "正文"
HEADING_STYLE_2 = "标题 2"
HEADING_STYLE_3 = "标题 3"
HEADING_STYLE_4 = "标题 4"


def compose_toc(structure: dict, file_path: str=None):
    # generate table of contents in a docx
    if os.path.exists(file_path):
        doc = Document(file_path)
    else:
        doc = Document()
    
    _modify_toc_styles(doc)

    title = doc.add_paragraph("目录")
    title.runs[0].font.name = "Times New Roman"
    title.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title.runs[0].font.size = Pt(14)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for chapter in structure["chapters"]:
        # 1-level
        doc.add_paragraph(chapter["title"], style=TOC_STYLE_1)

        # 2-level
        if "sections" in chapter:
            for section in chapter["sections"]:
                doc.add_paragraph(section["title"], style=TOC_STYLE_2)

                # 3-level 
                if "subsections" in section:
                    for subsection in section["subsections"]:
                        doc.add_paragraph(subsection["title"], style=TOC_STYLE_3) 
    doc.add_page_break()
    doc.save(file_path)  

def compose_main_body(title: str, structure: dict, file_path: str=None):
    
    doc = Document(file_path)

    _modify_normal_style(doc)
    _modify_heading_styles(doc)

    # generate main body of docx
    # tasks = []
    # for chapter in structure["chapters"]:
    #     # 1-level
    #     if "sections" not in chapter:
    #         tasks.append(chapter["title"])

    #     # 2-level
    #     else:
    #         for section in chapter["sections"]:
    #             if "subsections" not in section:
    #                 tasks.append(f"{chapter['title']} -> {section['title']}")

    #             # 3-level 
    #             else:
    #                 for subsection in section["subsections"]:
    #                     tasks.append(f"{chapter['title']} -> {section['title']} -> {subsection['title']}")
    

    for chapter in structure["chapters"]:
        # 1-level Handling
        doc.add_paragraph(chapter["title"], style=HEADING_STYLE_2)

        if "sections" not in chapter:
            # Main Body     
            texts = generate_section_text(title=title, section=chapter["title"], structure=structure)
            for text in texts:
                doc.add_paragraph(text, style=NORMAL_STYLE) 

        # 2-level
        else:
            for section in chapter["sections"]:
                # 2-level Handling
                doc.add_paragraph(section['title'], style=HEADING_STYLE_3)

                if "subsections" not in section:
                    # Main Body
                    texts = generate_section_text(title=title, section=f"{chapter['title']} -> {section['title']}", structure=structure)
                    for text in texts:
                        doc.add_paragraph(text, style=NORMAL_STYLE)

                # 3-level 
                else:
                    for subsection in section["subsections"]:
                        # 3-level Handling
                        doc.add_paragraph(subsection['title'], style=HEADING_STYLE_4)
                        # Main Body
                        texts = generate_section_text(title=title, section=f"{chapter['title']} -> {section['title']} -> {subsection['title']}", structure=structure)
                        for text in texts:
                            doc.add_paragraph(text, style=NORMAL_STYLE)

        doc.add_page_break()   

    doc.save(file_path)


def _modify_toc_styles(doc: Document):
    if TOC_STYLE_1 in doc.styles:
        style = doc.styles[TOC_STYLE_1]
    else:
        style = doc.styles.add_style(TOC_STYLE_1, WD_STYLE_TYPE.PARAGRAPH)

    style.font.size = Pt(12)  # 12磅（小四）
    style.font.bold = True
    style.font.name = "Times New Roman"
    style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    if TOC_STYLE_2 in doc.styles:
        style = doc.styles[TOC_STYLE_2]
    else:
        style = doc.styles.add_style(TOC_STYLE_2, WD_STYLE_TYPE.PARAGRAPH)

    style.font.size = Pt(12)  # 12磅（小四）
    style.font.bold = False
    style.font.name = "Times New Roman"
    style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.left_indent = Pt(20)

    if TOC_STYLE_3 in doc.styles:
        style = doc.styles[TOC_STYLE_3]
    else:
        style = doc.styles.add_style(TOC_STYLE_3, WD_STYLE_TYPE.PARAGRAPH)

    style.font.size = Pt(12)  # 12磅（小四）
    style.font.bold = False
    style.font.name = "Times New Roman"
    style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.left_indent = Pt(40)




# def _set_toc_style(paragraph: Paragraph, level: int):
#     """
#         目录（TOC）：段前段后0，行间距1.25
#     """
#     run = paragraph.runs[0]

#     run.font.size = Pt(12)  # 12磅（小四）
#     run.font.name = "Times New Roman"
#     paragraph.paragraph_format.line_spacing = 1.25
#     paragraph.paragraph_format.space_before = Pt(0)
#     paragraph.paragraph_format.space_after = Pt(0)

#     if level == 1:
#         run.font.bold = True
#         run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
#     elif level == 2:
#         run.font.bold = False
#         run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#         paragraph.paragraph_format.left_indent = Pt(20)
#     elif level == 3:
#         run.font.bold = False
#         run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#         paragraph.paragraph_format.left_indent = Pt(40)


def _modify_normal_style(doc: Document):

    if NORMAL_STYLE in doc.styles:
        style = doc.styles[NORMAL_STYLE]
    else:
        style = doc.styles.add_style(NORMAL_STYLE, WD_STYLE_TYPE.PARAGRAPH)
        
    # font style
    font_size = Pt(12)
    style.font.size = font_size
    style.font.bold = False
    style.font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # jusify alignment, first line indent 2 characters
    style.paragraph_format.first_line_indent = font_size * 2 
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def _modify_heading_styles(doc: Document):

    if HEADING_STYLE_2 in doc.styles:
        style = doc.styles[HEADING_STYLE_2]
    else:
        style = doc.styles.add_style(HEADING_STYLE_2, WD_STYLE_TYPE.PARAGRAPH)

    # font
    style.font.size = Pt(16)
    style.font.bold = True
    style.font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # para
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    style.paragraph_format.space_before = Pt(13)

    if HEADING_STYLE_3 in doc.styles:
        style = doc.styles[HEADING_STYLE_3]
    else:
        style = doc.styles.add_style(HEADING_STYLE_3, WD_STYLE_TYPE.PARAGRAPH)

    # font
    style.font.size = Pt(14)
    style.font.bold = True
    style.font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # para
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    style.paragraph_format.space_before = Pt(13)


    if HEADING_STYLE_4 in doc.styles:
        style = doc.styles[HEADING_STYLE_4]
    else:
        style = doc.styles.add_style(HEADING_STYLE_4, WD_STYLE_TYPE.PARAGRAPH)

    # font
    style.font.size = Pt(12)
    style.font.bold = True
    style.font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # para
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    style.paragraph_format.space_before = Pt(13)